"""DOCX parser for curriculum data extraction."""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .schemas import (
    AdmissionCriteriaImportData,
    BlockType,
    CourseGroupImportData,
    CourseImportData,
    CoursePLOContributionImportData,
    CurriculumImportData,
    KnowledgeBlockImportData,
    LearningOutcomeImportData,
    MappingStrength,
    ProgramCoursePlacementImportData,
    ProgramImportData,
    ProgramObjectiveImportData,
    ProgramVersionImportData,
    PLOPOMapImportData,
)

logger = logging.getLogger(__name__)


def normalize_text(text: str) -> str:
    """Normalize whitespace and strip artefacts from DOCX text."""
    if text is None:
        return ""
    cleaned = text.replace("\xa0", " ")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


ADMISSION_FIELD_LABELS = {
    "target_group": [
        "đối tượng tuyển sinh",
        "đối tượng",
        "admission target",
        "target intake",
    ],
    "selection_criteria": [
        "tiêu chí tuyển sinh",
        "tiêu chí",
        "admission requirement",
        "admission criteria",
    ],
}


def _remove_numeric_prefix(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*(?:[\.\)])?\s*", "", text)


def identify_admission_field(text: str) -> Optional[str]:
    cleaned = normalize_text(text)
    cleaned = _remove_numeric_prefix(cleaned)
    lowered = cleaned.lower()
    for field, labels in ADMISSION_FIELD_LABELS.items():
        for label in labels:
            if lowered.startswith(label):
                return field
    return None


def strip_admission_label(field_key: str, content: str) -> str:
    cleaned = normalize_text(content)
    cleaned = _remove_numeric_prefix(cleaned)
    lowered = cleaned.lower()
    for label in ADMISSION_FIELD_LABELS.get(field_key, []):
        if lowered.startswith(label):
            stripped = cleaned[len(label):].lstrip(":：-– ).").strip()
            return stripped
    return cleaned


def merge_admission_content(existing: Optional[str], addition: str) -> str:
    cleaned = normalize_text(addition)
    if not cleaned:
        return existing or ""
    if not existing:
        return cleaned
    existing_norm = existing.strip()
    if cleaned.lower() == existing_norm.lower():
        return existing_norm
    return f"{existing_norm}\n{cleaned}".strip()


def split_codes(value: str) -> List[str]:
    """Split comma/semicolon separated codes."""
    if not value:
        return []
    tokens = re.split(r"[;,/]| và | and ", value, flags=re.IGNORECASE)
    cleaned = [normalize_text(token).upper() for token in tokens if normalize_text(token)]
    return cleaned


_PI_FORMAT_PATTERN = re.compile(r"^PI\d+\.\d+$", re.IGNORECASE)


def normalize_performance_indicator_value(
    plo_code: str,
    raw_indicator: Optional[str],
    fallback_level: Optional[int] = None,
) -> str:
    """Normalise diverse PI notations into the canonical PIx.y format."""
    cleaned_indicator = normalize_text(raw_indicator or "")
    cleaned_indicator_upper = cleaned_indicator.upper()
    if cleaned_indicator_upper and _PI_FORMAT_PATTERN.match(cleaned_indicator_upper):
        numbers = re.findall(r"\d+", cleaned_indicator_upper)
        major = numbers[0] if numbers else "1"
        minor = numbers[1] if len(numbers) >= 2 else (str(fallback_level) if fallback_level is not None else "1")
        major = major.lstrip("0") or "0"
        minor = minor.lstrip("0") or "0"
        return f"PI{major}.{minor}"

    tokens = [
        token
        for token in re.split(r"[^\w]+", cleaned_indicator_upper)
        if token
    ]
    digits: List[str] = []
    for token in tokens:
        if token.startswith("PLO"):
            continue
        if token == "PI":
            continue
        if token.startswith("PI") and token[2:].isdigit():
            digits.append(token[2:])
            continue
        if token.isdigit():
            digits.append(token)

    plo_digits = re.findall(r"\d+", normalize_text(plo_code).upper())
    major = digits[0] if digits else (plo_digits[0] if plo_digits else None)
    minor = digits[1] if len(digits) >= 2 else None

    if not major and plo_digits:
        major = plo_digits[0]
    if not major:
        major = "1"

    if not minor:
        if fallback_level is not None:
            minor = str(fallback_level)
        elif digits:
            minor = digits[0]
        elif plo_digits:
            minor = plo_digits[0]
        else:
            minor = "1"

    if major.isdigit():
        major = str(int(major))
    if minor.isdigit():
        minor = str(int(minor))

    return f"PI{major}.{minor}"


def safe_int(value: str) -> Optional[int]:
    """Best-effort integer parsing from arbitrary strings."""
    match = re.search(r"\d+", value or "")
    if not match:
        return None
    try:
        return int(match.group())
    except ValueError:
        return None


def is_probably_table(headers: Iterable[str], keywords: Iterable[str]) -> bool:
    header_text = " ".join(normalize_text(h).lower() for h in headers)
    return any(keyword in header_text for keyword in keywords)


@dataclass
class ParserContext:
    """Mutable container that accumulates parsed artefacts prior to model conversion."""

    program_fields: Dict[str, Any] = field(default_factory=dict)
    version_fields: Dict[str, Any] = field(default_factory=dict)
    objectives_raw: List[Dict[str, Any]] = field(default_factory=list)
    outcomes_raw: List[Dict[str, Any]] = field(default_factory=list)
    plo_po_raw: List[Dict[str, Any]] = field(default_factory=list)
    admission_raw: Dict[str, Any] = field(default_factory=dict)
    knowledge_blocks_raw: List[Dict[str, Any]] = field(default_factory=list)
    knowledge_block_lookup: Dict[str, Dict[str, Any]] = field(default_factory=dict)
    course_groups_raw: List[Dict[str, Any]] = field(default_factory=list)
    course_group_lookup: Dict[str, Dict[str, Any]] = field(default_factory=dict)
    courses_raw: Dict[str, Dict[str, Any]] = field(default_factory=dict)
    placements_raw: List[Dict[str, Any]] = field(default_factory=list)
    course_plo_raw: List[Dict[str, Any]] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    blocking_errors: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    unassigned_courses: Set[str] = field(default_factory=set)
    placement_signatures: Set[Tuple[str, Optional[str], Optional[str], Optional[int]]] = field(default_factory=set)
    admission_active_field: Optional[str] = None

    def add_program_field(self, key: str, value: Any, source: str):
        if not value:
            return
        current = normalize_text(str(value))
        if not current:
            return
        existing = self.program_fields.get(key)
        if existing and normalize_text(existing) != current:
            self.warnings.append(
                f"Giá trị trường chương trình '{key}' không khớp (đang dùng '{existing}', nguồn mới '{current}' từ {source}). "
                "Giữ lại giá trị gốc."
            )
            return
        self.program_fields[key] = current

    def add_version_field(self, key: str, value: Any, source: str):
        if not value:
            return
        current = normalize_text(str(value))
        if not current:
            return
        existing = self.version_fields.get(key)
        if existing and normalize_text(existing) != current:
            self.warnings.append(
                f"Giá trị trường phiên bản '{key}' không khớp (đang dùng '{existing}', nguồn mới '{current}' từ {source}). "
                "Giữ lại giá trị gốc."
            )
            return
        self.version_fields[key] = current

    def add_objective(self, code: str, description: str, metadata: Optional[Dict[str, Any]] = None):
        code_norm = normalize_text(code).upper()
        if not code_norm:
            return
        for existing in self.objectives_raw:
            if existing["code"] == code_norm:
                if description and not existing["description"]:
                    existing["description"] = normalize_text(description)
                return
        self.objectives_raw.append(
            {
                "code": code_norm,
                "description": normalize_text(description),
                "metadata": metadata or {},
            }
        )

    def add_outcome(
        self,
        code: str,
        description: str,
        competency_level: Optional[float],
        metadata: Optional[Dict[str, Any]] = None,
    ):
        code_norm = normalize_text(code).upper()
        if not code_norm:
            return
        for existing in self.outcomes_raw:
            if existing["code"] == code_norm:
                if description and not existing["description"]:
                    existing["description"] = normalize_text(description)
                if competency_level is not None and existing["competency_level"] is None:
                    existing["competency_level"] = competency_level
                return
        self.outcomes_raw.append(
            {
                "code": code_norm,
                "description": normalize_text(description),
                "competency_level": competency_level,
                "metadata": metadata or {},
            }
        )

    def add_mapping(self, plo_code: str, po_code: str, strength: MappingStrength, source: str):
        plo_code_norm = normalize_text(plo_code).upper()
        po_code_norm = normalize_text(po_code).upper()
        if not plo_code_norm or not po_code_norm:
            return
        self.plo_po_raw.append(
            {
                "plo_code": plo_code_norm,
                "po_code": po_code_norm,
                "mapping_strength": strength,
                "metadata": {"source": source},
            }
        )

    def ensure_block(self, name: str, **kwargs) -> Dict[str, Any]:
        key = normalize_text(name)
        if not key:
            return {}
        if key in self.knowledge_block_lookup:
            block = self.knowledge_block_lookup[key]
            if kwargs.get("block_type") and not block.get("block_type"):
                block["block_type"] = kwargs["block_type"]
            return block
        block = {
            "name": normalize_text(name),
            "block_type": kwargs.get("block_type"),
            "parent_block_name": kwargs.get("parent_block_name"),
            "display_order": kwargs.get("display_order") or len(self.knowledge_blocks_raw) + 1,
            "description": kwargs.get("description"),
            "metadata": kwargs.get("metadata", {}),
        }
        self.knowledge_blocks_raw.append(block)
        self.knowledge_block_lookup[key] = block
        return block

    def ensure_course_group(self, name: str, block_name: Optional[str], **kwargs) -> Dict[str, Any]:
        key = normalize_text(name)
        if not key:
            return {}
        if key in self.course_group_lookup:
            group = self.course_group_lookup[key]
            if block_name and not group.get("block_name"):
                group["block_name"] = normalize_text(block_name)
            return group
        group = {
            "name": normalize_text(name),
            "block_name": normalize_text(block_name) if block_name else None,
            "group_type": kwargs.get("group_type"),
            "display_order": kwargs.get("display_order") or len(self.course_groups_raw) + 1,
            "min_credits_required": kwargs.get("min_credits_required", 0),
            "max_credits_allowed": kwargs.get("max_credits_allowed"),
            "description": kwargs.get("description"),
            "metadata": kwargs.get("metadata", {}),
        }
        self.course_groups_raw.append(group)
        self.course_group_lookup[key] = group
        return group

    def upsert_course(self, course: Dict[str, Any]):
        code = normalize_text(course.get("code", "")).upper()
        if not code:
            return
        existing = self.courses_raw.get(code)
        if existing:
            for field in ["name", "en_name", "default_total_credits"]:
                if course.get(field) and not existing.get(field):
                    existing[field] = course[field]
            return
        self.courses_raw[code] = course

    def add_course_placement(self, placement: Dict[str, Any]):
        if not placement.get("course_code"):
            return
        placement["course_code"] = normalize_text(placement["course_code"]).upper()
        if placement.get("group_name"):
            placement["group_name"] = normalize_text(placement["group_name"])
        if placement.get("block_name"):
            placement["block_name"] = normalize_text(placement["block_name"])
        if not placement.get("group_name") and not placement.get("block_name"):
            placeholder_block = "Chưa phân nhóm"
            if placement["course_code"] not in self.unassigned_courses:
                self.unassigned_courses.add(placement["course_code"])
            self.ensure_block(placeholder_block, metadata={"source": "auto_generated"})
            placement["block_name"] = placeholder_block
        signature = (
            placement["course_code"],
            placement.get("group_name"),
            placement.get("block_name"),
            placement.get("semester_recommended"),
        )
        if signature in self.placement_signatures:
            return
        self.placement_signatures.add(signature)
        self.placements_raw.append(placement)

    def add_course_plo(self, contribution: Dict[str, Any]):
        course_code = normalize_text(contribution.get("course_code", "")).upper()
        plo_code = normalize_text(contribution.get("plo_code", "")).upper()
        if not course_code or not plo_code:
            return
        contribution["course_code"] = course_code
        contribution["plo_code"] = plo_code
        contribution_level = contribution.get("contribution_level")
        contribution["performance_indicator"] = normalize_performance_indicator_value(
            plo_code,
            contribution.get("performance_indicator"),
            contribution_level,
        )
        self.course_plo_raw.append(contribution)

    def add_warning(self, message: str):
        if message not in self.warnings:
            self.warnings.append(message)

    def add_blocking_error(self, message: str):
        if message not in self.blocking_errors:
            self.blocking_errors.append(message)

    def build_curriculum(self) -> Optional[CurriculumImportData]:
        """Convert accumulated raw data into Pydantic models."""
        missing_program_fields = [
            name
            for name in ("code", "vn_name", "en_name", "faculty_name")
            if not self.program_fields.get(name)
        ]
        if missing_program_fields:
            self.add_blocking_error(
                "Thiếu thông tin chương trình trong tài liệu. Cần bổ sung: "
                + ", ".join(missing_program_fields)
            )
            return None

        # Normalise program fields
        program_code = normalize_text(self.program_fields["code"]).upper()
        self.program_fields["code"] = re.sub(r"\s+", "", program_code)
        self.program_fields["vn_name"] = normalize_text(self.program_fields["vn_name"])
        self.program_fields["en_name"] = normalize_text(
            self.program_fields.get("en_name") or self.program_fields["vn_name"]
        )
        self.program_fields["faculty_name"] = normalize_text(self.program_fields["faculty_name"])

        program_metadata = {
            "source": "docx_parser",
            "raw": {k: v for k, v in self.program_fields.items()},
        }

        program = ProgramImportData(
            code=self.program_fields["code"],
            vn_name=self.program_fields["vn_name"],
            en_name=self.program_fields["en_name"],
            faculty_name=self.program_fields["faculty_name"],
            awarding_institution=self.program_fields.get("awarding_institution"),
            degree_title=self.program_fields.get("degree_title"),
            degree_level=self.program_fields.get("degree_level"),
            study_mode=self.program_fields.get("study_mode"),
            metadata=program_metadata,
        )

        version_number = self.version_fields.get("version_number")
        if not version_number:
            version_number = f"{datetime.now().year}-1"

        version_metadata = {
            "source": "docx_parser",
            "raw": {k: v for k, v in self.version_fields.items()},
        }

        total_credits_value = self.version_fields.get("total_credits")
        parsed_credits = safe_int(total_credits_value) if total_credits_value else None
        if parsed_credits is not None:
            total_credits_value = str(parsed_credits)

        program_version = ProgramVersionImportData(
            version_number=version_number,
            version_name=self.version_fields.get("version_name"),
            total_credits=total_credits_value,
            duration=self.version_fields.get("duration"),
            status=self.version_fields.get("status", "draft"),
            change_type=self.version_fields.get("change_type"),
            metadata=version_metadata,
        )

        objectives = []
        for idx, row in enumerate(self.objectives_raw, start=1):
            objectives.append(
                ProgramObjectiveImportData(
                    code=row["code"],
                    description=row["description"],
                    display_order=idx,
                    metadata=row.get("metadata") or {},
                )
            )

        outcomes = []
        for idx, row in enumerate(self.outcomes_raw, start=1):
            outcomes.append(
                LearningOutcomeImportData(
                    code=row["code"],
                    description=row["description"],
                    competency_level=row.get("competency_level"),
                    display_order=idx,
                    metadata=row.get("metadata") or {},
                )
            )

        mappings = [
            PLOPOMapImportData(**mapping) for mapping in self.plo_po_raw
        ]

        admission = None
        target_group = (self.admission_raw.get("target_group") or "").strip()
        selection_criteria = (self.admission_raw.get("selection_criteria") or "").strip()
        if target_group or selection_criteria:
            admission = AdmissionCriteriaImportData(
                target_group=target_group,
                selection_criteria=selection_criteria,
                metadata=self.admission_raw.get("metadata") or {},
            )

        knowledge_blocks = []
        for block in self.knowledge_blocks_raw:
            block_type = block.get("block_type")
            block_type_enum = None
            if isinstance(block_type, str):
                normalized = normalize_text(block_type).lower()
                mapping = {
                    "đại cương": BlockType.GENERAL,
                    "general": BlockType.GENERAL,
                    "cơ sở": BlockType.CORE,
                    "core": BlockType.CORE,
                    "chuyên ngành": BlockType.SPECIALIZED,
                    "specialized": BlockType.SPECIALIZED,
                    "tự chọn": BlockType.ELECTIVE,
                    "elective": BlockType.ELECTIVE,
                    "cơ bản": BlockType.FOUNDATION,
                    "foundation": BlockType.FOUNDATION,
                }
                block_type_enum = mapping.get(normalized)
            knowledge_blocks.append(
                KnowledgeBlockImportData(
                    name=block["name"],
                    block_type=block_type_enum,
                    parent_block_name=block.get("parent_block_name"),
                    display_order=block.get("display_order", len(knowledge_blocks) + 1),
                    description=block.get("description"),
                    metadata=block.get("metadata") or {},
                )
            )

        course_groups = [
            CourseGroupImportData(
                name=group["name"],
                block_name=group.get("block_name"),
                group_type=None,
                display_order=group.get("display_order", idx),
                min_credits_required=group.get("min_credits_required", 0),
                max_credits_allowed=group.get("max_credits_allowed"),
                description=group.get("description"),
                metadata=group.get("metadata") or {},
            )
            for idx, group in enumerate(self.course_groups_raw, start=1)
        ]

        courses = [
            CourseImportData(
                code=code,
                name=value.get("name", code),
                en_name=value.get("en_name"),
                default_total_credits=value.get("default_total_credits", 0),
                default_lecture_credits=value.get("default_lecture_credits", 0),
                default_practical_credits=value.get("default_practical_credits", 0),
                default_project_credits=value.get("default_project_credits", 0),
                default_internship_credits=value.get("default_internship_credits", 0),
                syllabus_summary=value.get("syllabus_summary"),
                metadata=value.get("metadata") or {},
            )
            for code, value in self.courses_raw.items()
        ]

        placements = [
            ProgramCoursePlacementImportData(
                course_code=placement["course_code"],
                group_name=placement.get("group_name"),
                block_name=placement.get("block_name"),
                total_credits=placement.get("total_credits"),
                lecture_credits=placement.get("lecture_credits"),
                practical_credits=placement.get("practical_credits"),
                project_credits=placement.get("project_credits"),
                internship_credits=placement.get("internship_credits"),
                prereq_course_codes=placement.get("prereq_course_codes", []),
                coreq_course_codes=placement.get("coreq_course_codes", []),
                prereq_condition=placement.get("prereq_condition"),
                semester_recommended=placement.get("semester_recommended"),
                year_recommended=placement.get("year_recommended"),
                is_mandatory=placement.get("is_mandatory", True),
                note=placement.get("note"),
                display_order=placement.get("display_order", idx),
                metadata=placement.get("metadata") or {},
            )
            for idx, placement in enumerate(self.placements_raw, start=1)
        ]

        if self.unassigned_courses:
            sampled = sorted(self.unassigned_courses)[:5]
            remaining = len(self.unassigned_courses) - len(sampled)
            sample_text = ", ".join(sampled)
            if remaining > 0:
                sample_text = f"{sample_text}, ... (+{remaining} more)"
            self.add_warning(
                "Các học phần chưa xác định khối/nhóm đã được đưa vào khối 'Chưa phân nhóm': "
                + sample_text
            )

        course_plo_contributions: List[CoursePLOContributionImportData] = []
        for item in self.course_plo_raw:
            level = item.get("contribution_level")
            if not level:
                self.add_warning(
                    f"Bỏ qua liên kết học phần {item.get('course_code')} với PLO {item.get('plo_code')} "
                    "vì thiếu mức đóng góp hợp lệ."
                )
                continue
            normalized_indicator = normalize_performance_indicator_value(
                item["plo_code"],
                item.get("performance_indicator"),
                level,
            )
            course_plo_contributions.append(
                CoursePLOContributionImportData(
                    course_code=item["course_code"],
                    plo_code=item["plo_code"],
                    contribution_level=level,
                    performance_indicator=normalized_indicator,
                    note=item.get("note"),
                    metadata=item.get("metadata") or {},
                )
            )

        parse_metadata = {
            "parsed_at": datetime.now().isoformat(),
            "warnings_count": len(self.warnings),
            "tables_processed": self.metadata.get("table_stats", {}),
            "paragraphs_processed": self.metadata.get("paragraph_count", 0),
        }

        return CurriculumImportData(
            program=program,
            program_version=program_version,
            program_objectives=objectives,
            learning_outcomes=outcomes,
            plo_po_map=mappings,
            admission_criteria=admission,
            knowledge_blocks=knowledge_blocks,
            course_groups=course_groups,
            courses=courses,
            program_course_placement=placements,
            course_plo_contribution=course_plo_contributions,
            import_metadata=parse_metadata,
        )


class TableMapper:
    """Base class for table parsers with a priority indicator."""

    name = "base"
    priority = 100

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:  # pragma: no cover - interface
        raise NotImplementedError

    def process(self, table: Table, context: ParserContext):  # pragma: no cover - interface
        raise NotImplementedError


class ProgramInfoTableMapper(TableMapper):
    name = "program_info"
    priority = 10

    KEYWORDS = [
        "thông tin chung",
        "tên ngành",
        "program name",
        "mã ngành",
        "đơn vị quản lý",
        "khoa",
    ]

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:
        label_candidates = []
        for row in table.rows[:6]:
            cells = row.cells
            if len(cells) >= 2:
                label_candidates.append(normalize_text(cells[1].text).lower())
        score = sum(any(keyword in text for keyword in self.KEYWORDS) for text in label_candidates)
        return score >= 2

    def process(self, table: Table, context: ParserContext):
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if len(cells) == 0:
                continue

            # Scenario: "Tên ngành đào tạo" in second column, actual value third column.
            if len(cells) >= 3 and cells[1]:
                label = cells[1].lower()
                value = cells[2]
            else:
                # look for "Label: Value" in any cell
                joined = " ".join(cells)
                if ":" in joined:
                    label, value = joined.split(":", 1)
                elif "：" in joined:
                    label, value = joined.split("：", 1)
                else:
                    continue
                label = label.lower()
                value = value

            label_norm = normalize_text(label)
            value_norm = normalize_text(value)
            if not label_norm or not value_norm:
                continue

            pairs = []
            if ":" in value_norm:
                nested_label, nested_value = value_norm.split(":", 1)
                pairs.append((normalize_text(nested_label), normalize_text(nested_value)))
            pairs.append((label_norm, value_norm))

            for current_label, current_value in pairs:
                if not current_label or not current_value:
                    continue
                current_label_lower = current_label.lower()
                admission_field = identify_admission_field(current_label)

                if "mã ngành" in current_label_lower or "mã chương trình" in current_label_lower:
                    context.add_program_field("code", current_value, "program_info_table")
                elif "tên ngành" in current_label_lower or "tên chương trình" in current_label_lower or "tên tiếng việt" in current_label_lower:
                    context.add_program_field("vn_name", current_value, "program_info_table")
                elif "program name" in current_label_lower or "tên tiếng anh" in current_label_lower:
                    context.add_program_field("en_name", current_value, "program_info_table")
                elif "đơn vị quản lý" in current_label_lower or "khoa" in current_label_lower or "faculty" in current_label_lower:
                    context.add_program_field("faculty_name", current_value, "program_info_table")
                elif "trường cấp bằng" in current_label_lower or "awarding" in current_label_lower:
                    context.add_program_field("awarding_institution", current_value, "program_info_table")
                elif "tên gọi văn bằng" in current_label_lower or "degree title" in current_label_lower:
                    context.add_program_field("degree_title", current_value, "program_info_table")
                elif "trình độ đào tạo" in current_label_lower or "degree level" in current_label_lower:
                    context.add_program_field("degree_level", current_value, "program_info_table")
                elif "hình thức đào tạo" in current_label_lower or "study mode" in current_label_lower:
                    context.add_program_field("study_mode", current_value, "program_info_table")
                elif "số tín chỉ" in current_label_lower or "tổng số tín chỉ" in current_label_lower:
                    context.add_version_field("total_credits", current_value, "program_info_table")
                elif "thời gian đào tạo" in current_label_lower:
                    context.add_version_field("duration", current_value, "program_info_table")
                elif "khung" in current_label_lower or "phiên bản" in current_label_lower or "version name" in current_label_lower:
                    context.add_version_field("version_name", current_value, "program_info_table")
                elif "cập nhật" in current_label_lower:
                    context.add_version_field("status", "draft", "program_info_table")
                elif admission_field == "target_group":
                    existing = context.admission_raw.get("target_group")
                    context.admission_raw["target_group"] = (
                        merge_admission_content(existing, current_value) if existing else current_value
                    )
                    metadata = context.admission_raw.setdefault("metadata", {})
                    if metadata.get("source") not in ("admission_section_11", "admission_table"):
                        metadata["source"] = "program_info_table"
                elif admission_field == "selection_criteria":
                    existing = context.admission_raw.get("selection_criteria")
                    context.admission_raw["selection_criteria"] = (
                        merge_admission_content(existing, current_value) if existing else current_value
                    )
                    metadata = context.admission_raw.setdefault("metadata", {})
                    if metadata.get("source") not in ("admission_section_11", "admission_table"):
                        metadata["source"] = "program_info_table"

        if not context.version_fields.get("status"):
            context.version_fields["status"] = "draft"


class AdmissionTableMapper(TableMapper):
    name = "admission"
    priority = 40

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:
        sniff = " ".join(normalize_text(cell.text).lower() for row in table.rows[:2] for cell in row.cells)
        return "đối tượng" in sniff or "admission" in sniff

    def process(self, table: Table, context: ParserContext):
        target_value: Optional[str] = None
        selection_value: Optional[str] = None
        active_field: Optional[str] = None

        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if not any(cells):
                continue

            matched_field = None
            for idx, cell_text in enumerate(cells):
                field_key = identify_admission_field(cell_text)
                if not field_key:
                    continue
                matched_field = field_key
                active_field = field_key

                value_segments: List[str] = []
                stripped = strip_admission_label(field_key, cell_text)
                if stripped:
                    value_segments.append(stripped)
                value_segments.extend(cells[idx + 1:])

                combined = self._collect_value(value_segments)
                if combined:
                    if field_key == "target_group":
                        target_value = merge_admission_content(target_value, combined)
                    else:
                        selection_value = merge_admission_content(selection_value, combined)
                break

            if matched_field:
                continue

            if active_field:
                combined = self._collect_value(cells)
                if combined:
                    if active_field == "target_group":
                        target_value = merge_admission_content(target_value, combined)
                    else:
                        selection_value = merge_admission_content(selection_value, combined)

        if target_value is not None:
            context.admission_raw["target_group"] = target_value
        if selection_value is not None:
            context.admission_raw["selection_criteria"] = selection_value
        if context.admission_raw:
            context.admission_raw.setdefault("metadata", {})["source"] = "admission_table"

    @staticmethod
    def _collect_value(segments: Iterable[str]) -> str:
        value: Optional[str] = None
        for segment in segments:
            cleaned = normalize_text(segment)
            if not cleaned:
                continue
            value = merge_admission_content(value, cleaned)
        return value or ""


class PloDefinitionTableMapper(TableMapper):
    name = "plo_definition"
    priority = 20

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:
        header_text = " ".join(normalize_text(header).lower() for header in headers)
        return "chuẩn đầu ra" in header_text and ("tương ứng" in header_text or "muc tieu" in header_text)

    def process(self, table: Table, context: ParserContext):
        for row in table.rows[1:]:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if not cells:
                continue
            code = cells[0]
            if not re.match(r"^plo\d+", code, re.IGNORECASE):
                continue
            description = cells[1] if len(cells) > 1 else ""
            mapping_cell = cells[2] if len(cells) > 2 else ""
            competency_cell = cells[3] if len(cells) > 3 else ""

            competency_level = None
            if competency_cell:
                try:
                    competency_level = float(competency_cell.replace(",", "."))
                except ValueError:
                    competency_level = None

            context.add_outcome(
                code,
                description,
                competency_level,
                metadata={"source": "plo_definition_table"}
            )

            po_codes = [normalize_text(item).upper() for item in re.split(r"[,;]", mapping_cell) if normalize_text(item)]
            for po_code in po_codes:
                if not re.match(r"^PO\d+", po_code, re.IGNORECASE):
                    continue
                context.add_mapping(code, po_code, MappingStrength.STRONG, "plo_definition_table")


class PloPoMatrixTableMapper(TableMapper):
    name = "plo_po_matrix"
    priority = 30

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:
        combined_rows = [normalize_text(cell.text).lower() for row in table.rows[:2] for cell in row.cells]
        header_text = " ".join(combined_rows)
        return "plo" in header_text and "po" in header_text

    def process(self, table: Table, context: ParserContext):
        header_row_index = None
        plo_columns: List[Tuple[int, str]] = []

        for idx, row in enumerate(table.rows):
            cells = [normalize_text(cell.text) for cell in row.cells]
            if any(re.match(r"^plo\d+", cell, re.IGNORECASE) for cell in cells):
                header_row_index = idx
                for col_idx, cell in enumerate(cells):
                    if re.match(r"^plo\d+", cell, re.IGNORECASE):
                        plo_columns.append((col_idx, cell))
                break

        if header_row_index is None or not plo_columns:
            context.add_warning("Không thể xác định tiêu đề PLO trong ma trận PLO-PO.")
            return

        for row in table.rows[header_row_index + 1:]:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if not cells:
                continue
            po_code = cells[0]
            if not re.match(r"^po\d+", po_code, re.IGNORECASE):
                continue
            for col_idx, plo_code in plo_columns:
                if col_idx >= len(cells):
                    continue
                raw_value = normalize_text(cells[col_idx])
                strength = self._interpret_strength(raw_value)
                if strength:
                    context.add_mapping(plo_code, po_code, strength, "plo_po_matrix")
                elif raw_value:
                    context.add_warning(
                        f"Không thể diễn giải giá trị '{raw_value}' trong ô ma trận ({plo_code}, {po_code}). "
                        "Liên kết giữa PLO và PO này sẽ được bỏ qua."
                    )

    @staticmethod
    def _interpret_strength(value: str) -> Optional[MappingStrength]:
        if not value:
            return None
        lower = value.lower()
        if lower in {"1", "x", "✓", "✔", "●", "c", "strong", "cao", "high"}:
            return MappingStrength.STRONG
        if lower in {"2", "○", "b", "moderate", "medium", "trung bình"}:
            return MappingStrength.MODERATE
        if lower in {"3", "△", "a", "weak", "low", "thấp"}:
            return MappingStrength.WEAK
        if lower.isdigit():
            num = int(lower)
            if num >= 3:
                return MappingStrength.STRONG
            if num == 2:
                return MappingStrength.MODERATE
            if num == 1:
                return MappingStrength.WEAK
        return None


class KnowledgeStructureTableMapper(TableMapper):
    name = "knowledge_structure"
    priority = 50

    BLOCK_KEYWORDS = ["khối kiến thức", "block", "kiến thức"]
    GROUP_KEYWORDS = ["nhóm kiến thức", "course group", "nhóm học phần"]
    CREDIT_KEYWORDS = ["tín chỉ", "credit"]

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:
        return is_probably_table(headers, self.BLOCK_KEYWORDS)

    def process(self, table: Table, context: ParserContext):
        ignored_tables = context.metadata.setdefault("ignored_tables", [])
        if "knowledge_structure" not in ignored_tables:
            ignored_tables.append("knowledge_structure")

    @staticmethod
    def _index(headers: List[str], keywords: List[str]) -> Optional[int]:
        for idx, header in enumerate(headers):
            if any(keyword in header for keyword in keywords):
                return idx
        return None


class CourseListTableMapper(TableMapper):
    name = "course_list"
    priority = 60

    MAIN_BLOCK_REGEX = re.compile(r"^(?P<roman>[IVXLCDM]+)\.\s+(?P<name>.+)$")
    SUB_BLOCK_REGEX = re.compile(r"^(?P<roman>[IVXLCDM]+)\.(?P<number>\d+)\.?\s*(?P<name>.*)$")
    GROUP_REGEX = re.compile(r"^Nhóm\s*(?P<number>\d+)(?::\s*(?P<name>.*))?$", re.IGNORECASE)

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:
        header_text = " ".join(normalize_text(h).lower() for h in headers)
        return any(keyword in header_text for keyword in ["mã học phần", "mã số hp", "course code"])

    def process(self, table: Table, context: ParserContext):
        headers = [normalize_text(cell.text).lower() for cell in table.rows[0].cells]
        code_idx = self._find(headers, ["mã học phần", "mã số hp", "course code", "mã hp", "code"])
        name_idx = self._find(headers, ["tên học phần", "course name", "tên hp", "title"])
        credits_idx = self._find(headers, ["số tín chỉ", "credits", "tc"])
        lecture_idx = self._find(headers, ["lt", "lecture"])
        practical_idx = self._find(headers, ["th", "practical"])
        project_idx = self._find(headers, ["da", "project"])
        internship_idx = self._find(headers, ["thực tập", "internship"])
        semester_idx = self._find(headers, ["học kỳ", "semester"])
        prereq_idx = self._find(headers, ["học trước", "prereq"])
        coreq_idx = self._find(headers, ["song hành", "coreq"])

        block_display_order = 0
        sub_block_display_order: Dict[str, int] = {}
        group_display_order: Dict[str, int] = {}
        blocks_by_roman: Dict[str, str] = {}

        current_block_name: Optional[str] = None
        current_sub_block_name: Optional[str] = None
        current_group_name: Optional[str] = None

        for row_index, row in enumerate(table.rows):
            cells = [normalize_text(cell.text) for cell in row.cells]
            if not any(cells):
                continue

            if row_index == 0:
                continue

            primary_text = self._first_non_empty(cells)

            main_block = self._parse_main_block(primary_text)
            if main_block:
                block_name, roman = main_block
                block_key = normalize_text(block_name)
                existing = context.knowledge_block_lookup.get(block_key)
                if not existing:
                    block_display_order += 1
                display_order = existing.get("display_order") if existing else block_display_order
                context.ensure_block(
                    block_name,
                    display_order=display_order,
                    metadata={"source": "detailed_curriculum_table", "row_index": row_index + 1},
                )
                blocks_by_roman[roman] = block_name
                current_block_name = block_name
                current_sub_block_name = None
                current_group_name = None
                continue

            if self._looks_like_uppercase_block(primary_text):
                block_name = primary_text
                block_key = normalize_text(block_name)
                existing = context.knowledge_block_lookup.get(block_key)
                if not existing:
                    block_display_order += 1
                display_order = existing.get("display_order") if existing else block_display_order
                context.ensure_block(
                    block_name,
                    display_order=display_order,
                    metadata={"source": "detailed_curriculum_table", "row_index": row_index + 1},
                )
                current_block_name = block_name
                current_sub_block_name = None
                current_group_name = None
                continue

            sub_block = self._parse_sub_block(primary_text)
            if sub_block:
                roman, block_name = sub_block
                parent_block = blocks_by_roman.get(roman) or current_block_name
                if not parent_block:
                    auto_parent_name = f"Khối {roman}"
                    parent_key = normalize_text(auto_parent_name)
                    existing_parent = context.knowledge_block_lookup.get(parent_key)
                    if not existing_parent:
                        block_display_order += 1
                    display_order_parent = existing_parent.get("display_order") if existing_parent else block_display_order
                    context.ensure_block(
                        auto_parent_name,
                        display_order=display_order_parent,
                        metadata={
                            "source": "detailed_curriculum_table",
                            "row_index": row_index + 1,
                            "generated": True,
                        },
                    )
                    parent_block = auto_parent_name
                blocks_by_roman.setdefault(roman, parent_block)

                block_key = normalize_text(block_name)
                existing_sub = context.knowledge_block_lookup.get(block_key)
                if not existing_sub:
                    sub_block_display_order[parent_block] = sub_block_display_order.get(parent_block, 0) + 1
                display_order = existing_sub.get("display_order") if existing_sub else sub_block_display_order[parent_block]

                context.ensure_block(
                    block_name,
                    parent_block_name=parent_block,
                    display_order=display_order,
                    metadata={"source": "detailed_curriculum_table", "row_index": row_index + 1},
                )

                current_block_name = parent_block
                current_sub_block_name = block_name
                current_group_name = None
                continue

            group_info = self._parse_group(primary_text)
            if group_info:
                group_name, generated = group_info
                parent_for_group = current_sub_block_name or current_block_name
                if not parent_for_group:
                    context.add_warning(
                        f"Không thể gán nhóm kiến thức '{group_name}' vì thiếu khối cha (dòng {row_index + 1})."
                    )
                    current_group_name = None
                    continue

                existing_group = context.course_group_lookup.get(normalize_text(group_name))
                if not existing_group:
                    group_display_order[parent_for_group] = group_display_order.get(parent_for_group, 0) + 1
                display_order = (
                    existing_group.get("display_order")
                    if existing_group
                    else group_display_order[parent_for_group]
                )

                metadata = {
                    "source": "detailed_curriculum_table",
                    "row_index": row_index + 1,
                }
                if generated:
                    metadata["generated_name"] = True

                context.ensure_course_group(
                    group_name,
                    parent_for_group,
                    display_order=display_order,
                    metadata=metadata,
                )
                current_group_name = group_name
                continue

            code = cells[code_idx] if code_idx is not None and code_idx < len(cells) else ""
            if not code:
                continue
            if not re.match(r"^[A-Z]{2,}\d+", code):
                continue

            course_payload = {
                "code": code,
                "name": cells[name_idx] if name_idx is not None and name_idx < len(cells) else code,
                "default_total_credits": safe_int(cells[credits_idx]) or 0 if credits_idx is not None and credits_idx < len(cells) else 0,
                "default_lecture_credits": safe_int(cells[lecture_idx]) or 0 if lecture_idx is not None and lecture_idx < len(cells) else 0,
                "default_practical_credits": safe_int(cells[practical_idx]) or 0 if practical_idx is not None and practical_idx < len(cells) else 0,
                "default_project_credits": safe_int(cells[project_idx]) or 0 if project_idx is not None and project_idx < len(cells) else 0,
                "default_internship_credits": safe_int(cells[internship_idx]) or 0 if internship_idx is not None and internship_idx < len(cells) else 0,
                "metadata": {"source": "course_list_table"},
            }
            context.upsert_course(course_payload)

            placement_payload = {
                "course_code": code,
                "block_name": current_sub_block_name or current_block_name,
                "group_name": current_group_name,
                "total_credits": course_payload["default_total_credits"],
                "lecture_credits": course_payload["default_lecture_credits"],
                "practical_credits": course_payload["default_practical_credits"],
                "project_credits": course_payload["default_project_credits"],
                "internship_credits": course_payload["default_internship_credits"],
                "prereq_course_codes": split_codes(cells[prereq_idx]) if prereq_idx is not None and prereq_idx < len(cells) else [],
                "coreq_course_codes": split_codes(cells[coreq_idx]) if coreq_idx is not None and coreq_idx < len(cells) else [],
                "semester_recommended": safe_int(cells[semester_idx]) if semester_idx is not None and semester_idx < len(cells) else None,
                "metadata": {"source": "course_list_table"},
            }
            context.add_course_placement(placement_payload)

    @staticmethod
    def _find(headers: List[str], keywords: List[str]) -> Optional[int]:
        for idx, header in enumerate(headers):
            if any(keyword in header for keyword in keywords):
                return idx
        return None

    @staticmethod
    def _first_non_empty(cells: List[str]) -> str:
        for cell in cells:
            if cell:
                return cell
        return ""

    @staticmethod
    def _contains_letter(text: str) -> bool:
        return bool(re.search(r"[A-Za-zÀ-ỹ]", text))

    def _parse_main_block(self, text: str) -> Optional[Tuple[str, str]]:
        if not text:
            return None
        match = self.MAIN_BLOCK_REGEX.match(text)
        if not match:
            return None
        name = normalize_text(match.group("name"))
        if not self._contains_letter(name):
            return None
        block_name = f"{match.group('roman')}. {name}"
        return block_name, match.group("roman")

    def _parse_sub_block(self, text: str) -> Optional[Tuple[str, str]]:
        if not text:
            return None
        match = self.SUB_BLOCK_REGEX.match(text)
        if not match:
            return None
        name = normalize_text(match.group("name"))
        if not self._contains_letter(name):
            return None
        roman = match.group("roman")
        block_name = normalize_text(text)
        if not block_name:
            block_name = f"{roman}.{match.group('number')}"
        return roman, block_name

    def _parse_group(self, text: str) -> Optional[Tuple[str, bool]]:
        if not text:
            return None
        match = self.GROUP_REGEX.match(text)
        if not match:
            return None
        number = match.group("number")
        name_part = normalize_text(match.group("name") or "")
        if name_part:
            group_name = f"Nhóm {number}: {name_part}"
            generated = False
        else:
            group_name = f"Nhóm {number}"
            generated = True
        return group_name, generated

    @staticmethod
    def _looks_like_uppercase_block(text: str) -> bool:
        if not text:
            return False
        if len(text) < 4 or " " not in text:
            return False
        return text.isupper()


class CoursePloMatrixMapper(TableMapper):
    name = "course_plo_matrix"
    priority = 55

    def matches(self, headers: List[str], table: Table, context: ParserContext) -> bool:
        header_text = " ".join(normalize_text(h).lower() for h in headers)
        return "plo" in header_text and ("mã học phần" in header_text or "course" in header_text)

    def process(self, table: Table, context: ParserContext):
        plo_row_index = None
        pi_row_index = None
        plo_columns: Dict[int, str] = {}
        pi_lookup: Dict[int, str] = {}

        for idx, row in enumerate(table.rows[:3]):
            cells = [normalize_text(cell.text) for cell in row.cells]
            if any(re.match(r"^plo\d+", cell, re.IGNORECASE) for cell in cells):
                plo_row_index = idx
                for col_idx, cell in enumerate(cells):
                    if re.match(r"^plo\d+", cell, re.IGNORECASE):
                        plo_columns[col_idx] = cell
            if any(cell.lower().startswith("pi") for cell in cells):
                pi_row_index = idx

        if plo_row_index is None or not plo_columns:
            context.add_warning("Không thể xác định tiêu đề PLO trong ma trận học phần - PLO.")
            return

        if pi_row_index is not None:
            pi_row = [normalize_text(cell.text) for cell in table.rows[pi_row_index].cells]
            for col_idx, value in enumerate(pi_row):
                if col_idx in plo_columns and value:
                    pi_lookup[col_idx] = value

        data_start = max(plo_row_index, (pi_row_index or plo_row_index)) + 1

        for row in table.rows[data_start:]:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if not cells:
                continue
            course_code = cells[0]
            if not course_code or course_code.lower().startswith("mã học phần"):
                continue
            for col_idx, plo_code in plo_columns.items():
                if col_idx >= len(cells):
                    continue
                raw_value = cells[col_idx]
                level = self._map_contribution_level(raw_value)
                if level is not None:
                    pi_value = pi_lookup.get(col_idx)
                    performance_indicator = f"{plo_code}.{pi_value}" if pi_value else f"{plo_code}.{level}"
                    context.add_course_plo(
                        {
                            "course_code": course_code,
                            "plo_code": plo_code,
                            "contribution_level": level,
                            "performance_indicator": performance_indicator,
                            "metadata": {"source": "course_plo_matrix"},
                        }
                    )
                elif raw_value and raw_value not in {'-', '0'}:
                    context.add_warning(
                        f"Không thể diễn giải mức đóng góp '{raw_value}' cho học phần {course_code} và PLO {plo_code}."
                    )

    @staticmethod
    def _map_contribution_level(value: str) -> Optional[int]:
        if not value:
            return None
        lower = value.lower()
        mapping = {
            "h": 5,
            "high": 5,
            "m": 3,
            "medium": 3,
            "l": 1,
            "low": 1,
            "x": 5,
            "✓": 5,
            "✔": 5,
            "●": 5,
            "○": 3,
            "△": 1,
            "a": 1,
            "b": 3,
            "c": 5,
        }
        if lower in mapping:
            return mapping[lower]
        if lower.isdigit():
            num = int(lower)
            if 1 <= num <= 5:
                return num
        return None


class ParagraphExtractor:
    """Base class for paragraph-based heuristics."""

    def process(self, text: str, context: ParserContext):
        raise NotImplementedError


class ProgramParagraphExtractor(ParagraphExtractor):
    PATTERNS = {
        "code": re.compile(r"mã\s+(?:chương\s+trình|ngành)\s*[:：]?\s*([A-Z0-9\-]+)", re.IGNORECASE),
        "vn_name": re.compile(r"(?:tên\s+(?:ngành|chương\s+trình)\s*(?:tiếng\s+việt)?)\s*[:：]\s*(.+)", re.IGNORECASE),
        "en_name": re.compile(r"(?:tên\s+tiếng\s+anh|english\s+name)\s*[:：]\s*(.+)", re.IGNORECASE),
        "faculty_name": re.compile(r"(?:khoa|đơn\s+vị\s+quản\s+lý)\s*[:：]\s*(.+)", re.IGNORECASE),
    }

    def process(self, text: str, context: ParserContext):
        for key, pattern in self.PATTERNS.items():
            match = pattern.search(text)
            if match:
                context.add_program_field(key, match.group(1), "paragraph")


class ObjectiveParagraphExtractor(ParagraphExtractor):
    def process(self, text: str, context: ParserContext):
        match = re.match(r"^(PO\d+)\s*[:：\-–]\s*(.+)$", text, re.IGNORECASE)
        if match:
            context.add_objective(match.group(1), match.group(2))


class OutcomeParagraphExtractor(ParagraphExtractor):
    def process(self, text: str, context: ParserContext):
        match = re.match(r"^(PLO\d+)\s*[:：\-–]\s*(.+)$", text, re.IGNORECASE)
        if not match:
            return
        description = match.group(2)
        level_match = re.search(r"(?:TĐNL|level|cấp)\s*[:：]?\s*(\d(?:\.\d)?)", description, re.IGNORECASE)
        competency = float(level_match.group(1)) if level_match else None
        context.add_outcome(match.group(1), description, competency)


class AdmissionParagraphExtractor(ParagraphExtractor):
    SECTION_PATTERN = re.compile(r"^11\.\s*(chuẩn đầu vào|admission requirements?)", re.IGNORECASE)
    FIELD_PATTERN = re.compile(
        r"""
        ^11\.(?P<field>1|2)\s*
        (?:[\.:)\-–]\s*)?
        (?P<content>.*)$
        """,
        re.IGNORECASE | re.VERBOSE,
    )

    def process(self, text: str, context: ParserContext):
        normalized = normalize_text(text)
        if not normalized:
            return

        if self.SECTION_PATTERN.match(normalized):
            context.admission_active_field = None
            context.admission_raw.setdefault("metadata", {})["source"] = "admission_section_11"
            return

        field_match = self.FIELD_PATTERN.match(normalized)
        if field_match:
            field_key = "target_group" if field_match.group("field") == "1" else "selection_criteria"
            content = strip_admission_label(field_key, field_match.group("content") or "")
            context.admission_active_field = field_key
            if content:
                context.admission_raw[field_key] = merge_admission_content(context.admission_raw.get(field_key), content)
                context.admission_raw.setdefault("metadata", {})["source"] = "admission_section_11"
            else:
                # ensure metadata still tracks section 11 parsing even if content spans multiple paragraphs
                context.admission_raw.setdefault("metadata", {})["source"] = "admission_section_11"
            return

        if context.admission_active_field:
            if self._is_new_section(normalized):
                context.admission_active_field = None
            else:
                field_key = context.admission_active_field
                context.admission_raw[field_key] = merge_admission_content(
                    context.admission_raw.get(field_key), normalized
                )
                context.admission_raw.setdefault("metadata", {})["source"] = "admission_section_11"
                return

        lowered = normalized.lower()
        if "đối tượng tuyển sinh" in lowered or "target intake" in lowered:
            raw_value = normalized.split(":", 1)[-1].strip()
            value = strip_admission_label("target_group", raw_value)
            if value:
                existing = context.admission_raw.get("target_group")
                context.admission_raw["target_group"] = (
                    merge_admission_content(existing, value) if existing else value
                )
                metadata = context.admission_raw.setdefault("metadata", {})
                if metadata.get("source") != "admission_section_11":
                    metadata["source"] = "paragraph"
            context.admission_active_field = "target_group"
        elif "tiêu chí" in lowered or "criteria" in lowered:
            raw_value = normalized.split(":", 1)[-1].strip()
            value = strip_admission_label("selection_criteria", raw_value)
            if value:
                existing = context.admission_raw.get("selection_criteria")
                context.admission_raw["selection_criteria"] = (
                    merge_admission_content(existing, value) if existing else value
                )
                metadata = context.admission_raw.setdefault("metadata", {})
                if metadata.get("source") != "admission_section_11":
                    metadata["source"] = "paragraph"
            context.admission_active_field = "selection_criteria"

    @staticmethod
    def _is_new_section(text: str) -> bool:
        match = re.match(r"^(?P<main>\d+)(?:\.(?P<sub>\d+))?", text)
        if not match:
            return False
        main = int(match.group("main"))
        sub = match.group("sub")
        if main > 11:
            return True
        if main == 11 and sub and sub not in {"1", "2"}:
            return True
        return False


class DocxCurriculumParser:
    """Parser orchestrating paragraph and table extractors."""

    def __init__(self):
        self.context = ParserContext()
        self.table_mappers: List[TableMapper] = sorted(
            [
                ProgramInfoTableMapper(),
                AdmissionTableMapper(),
                PloDefinitionTableMapper(),
                PloPoMatrixTableMapper(),
                KnowledgeStructureTableMapper(),
                CourseListTableMapper(),
                CoursePloMatrixMapper(),
            ],
            key=lambda mapper: mapper.priority,
        )
        self.paragraph_extractors: List[ParagraphExtractor] = [
            ProgramParagraphExtractor(),
            ObjectiveParagraphExtractor(),
            OutcomeParagraphExtractor(),
            AdmissionParagraphExtractor(),
        ]
        self.current_section: Optional[str] = None

    def parse_from_path(self, file_path: str) -> Optional[CurriculumImportData]:
        try:
            path = Path(file_path)
            if not path.exists():
                self.context.add_blocking_error(f"Không tìm thấy tệp: {file_path}")
                return None
            if path.suffix.lower() != ".docx":
                self.context.add_blocking_error(
                    f"Định dạng tệp '{path.suffix}' không hợp lệ. Vui lòng sử dụng tệp .docx"
                )
                return None
            document = Document(file_path)
            return self._parse_document(document)
        except Exception as exc:
            logger.exception("Failed to parse DOCX from path %s", file_path)
            self.context.add_blocking_error(f"Không thể phân tích tệp DOCX: {exc}")
            return None

    def parse_from_bytes(self, file_bytes: bytes) -> Optional[CurriculumImportData]:
        try:
            document = Document(io.BytesIO(file_bytes))
            return self._parse_document(document)
        except Exception as exc:
            logger.exception("Failed to parse DOCX bytes")
            self.context.add_blocking_error(f"Không thể phân tích tệp DOCX: {exc}")
            return None

    def _parse_document(self, document: Document) -> Optional[CurriculumImportData]:
        paragraphs = document.paragraphs
        tables = document.tables

        self.context.metadata["paragraph_count"] = len(paragraphs)
        self.context.metadata["table_stats"] = {"total": len(tables)}

        self._process_paragraphs(paragraphs)
        self._process_tables(tables)

        curriculum = self.context.build_curriculum()
        if not curriculum and not self.context.blocking_errors:
            self.context.add_blocking_error("Không tìm thấy dữ liệu chương trình phù hợp trong tài liệu.")
        return curriculum

    def _process_paragraphs(self, paragraphs: List[Paragraph]):
        for paragraph in paragraphs:
            text = normalize_text(paragraph.text)
            if not text:
                continue
            if text.isupper() and len(text.split()) <= 6:
                # treat as section header
                self.current_section = text
                self.context.admission_active_field = None
                continue
            for extractor in self.paragraph_extractors:
                extractor.process(text, self.context)

    def _process_tables(self, tables: List[Table]):
        for index, table in enumerate(tables, start=1):
            if not table.rows:
                continue
            headers = [normalize_text(cell.text) for cell in table.rows[0].cells]
            matched_mappers = [mapper for mapper in self.table_mappers if mapper.matches(headers, table, self.context)]

            if len(matched_mappers) > 1:
                mapper_names = ", ".join(mapper.name for mapper in matched_mappers)
                self.context.add_warning(
                    f"Bảng {index} phù hợp với nhiều bộ phân tích ({mapper_names}). Sử dụng '{matched_mappers[0].name}'."
                )

            mapper = matched_mappers[0] if matched_mappers else None
            if not mapper:
                self.context.add_warning(f"Bảng {index} có cấu trúc không nhận diện được nên đã bị bỏ qua.")
                continue
            try:
                mapper.process(table, self.context)
            except Exception as exc:
                logger.exception("Error while processing table %s with mapper %s", index, mapper.name)
                self.context.add_warning(
                    f"Bảng {index} không thể xử lý bằng bộ phân tích '{mapper.name}': {exc}"
                )


def parse_docx_file(file_path: str) -> Tuple[Optional[CurriculumImportData], List[str], List[str]]:
    parser = DocxCurriculumParser()
    data = parser.parse_from_path(file_path)
    return data, parser.context.warnings, parser.context.blocking_errors


def parse_docx_bytes(file_bytes: bytes) -> Tuple[Optional[CurriculumImportData], List[str], List[str]]:
    parser = DocxCurriculumParser()
    data = parser.parse_from_bytes(file_bytes)
    return data, parser.context.warnings, parser.context.blocking_errors
